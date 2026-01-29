"""
Convert VIVA_PREPARATION_GUIDE.md to Word Document (.docx)
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def create_viva_docx():
    # Create a new Document
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    
    # Read the markdown file
    with open('VIVA_PREPARATION_GUIDE.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Parse and add content
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_content = []
    in_table = False
    table_rows = []
    
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_content = []
            else:
                in_code_block = False
                # Add code block as a paragraph with monospace font
                if code_content:
                    code_text = '\n'.join(code_content)
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        
        if in_code_block:
            code_content.append(line)
            i += 1
            continue
        
        # Handle tables
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            # Parse table row
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            # Skip separator rows (like |---|---|)
            if cells and not all(set(cell) <= {'-', ' ', ':'} for cell in cells):
                table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            # End of table
            in_table = False
            if table_rows:
                # Create table
                num_cols = len(table_rows[0]) if table_rows else 0
                if num_cols > 0:
                    table = doc.add_table(rows=len(table_rows), cols=num_cols)
                    table.style = 'Table Grid'
                    for row_idx, row_data in enumerate(table_rows):
                        for col_idx, cell_data in enumerate(row_data):
                            if col_idx < len(table.rows[row_idx].cells):
                                cell = table.rows[row_idx].cells[col_idx]
                                cell.text = cell_data
                                # Bold header row
                                if row_idx == 0:
                                    for run in cell.paragraphs[0].runs:
                                        run.bold = True
                    doc.add_paragraph()  # Add space after table
            table_rows = []
        
        # Handle headers
        if line.startswith('# '):
            # Main title
            title = line[2:].strip()
            # Remove emoji at start if present
            title = re.sub(r'^[^\w\s]+\s*', '', title)
            p = doc.add_heading(title, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            text = line[3:].strip()
            text = re.sub(r'^[^\w\s]+\s*', '', text)
            doc.add_heading(text, level=1)
        elif line.startswith('### '):
            text = line[4:].strip()
            text = re.sub(r'^[^\w\s]+\s*', '', text)
            doc.add_heading(text, level=2)
        elif line.startswith('> '):
            # Blockquote
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(text)
            run.italic = True
        elif line.startswith('- ') or line.startswith('* '):
            # Bullet point
            text = line[2:].strip()
            # Handle bold text
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            doc.add_paragraph(text, style='List Bullet')
        elif line.startswith('---'):
            # Horizontal rule - add a paragraph break
            doc.add_paragraph()
        elif line.strip():
            # Regular paragraph
            # Clean up markdown formatting
            text = line.strip()
            # Remove emoji codes
            text = re.sub(r':[a-z_]+:', '', text)
            # Handle bold text
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            # Handle inline code
            text = re.sub(r'`(.+?)`', r'\1', text)
            # Handle checkmarks and crosses
            text = text.replace('✓', '✓ ')
            text = text.replace('❌', '❌ ')
            text = text.replace('✅', '✅ ')
            
            if text:
                doc.add_paragraph(text)
        
        i += 1
    
    # Handle any remaining table
    if in_table and table_rows:
        num_cols = len(table_rows[0]) if table_rows else 0
        if num_cols > 0:
            table = doc.add_table(rows=len(table_rows), cols=num_cols)
            table.style = 'Table Grid'
            for row_idx, row_data in enumerate(table_rows):
                for col_idx, cell_data in enumerate(row_data):
                    if col_idx < len(table.rows[row_idx].cells):
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = cell_data
    
    # Save the document
    output_path = 'VIVA_PREPARATION_GUIDE.docx'
    doc.save(output_path)
    print(f"✅ Document saved successfully: {output_path}")
    return output_path

if __name__ == '__main__':
    create_viva_docx()
